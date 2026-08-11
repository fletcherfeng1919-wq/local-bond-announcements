"""Cross-check data/state_results.csv against provincial finance-bureau
(财政厅/财政局) 发行结果公告 pages -- a THIRD independent source alongside
celma.org.cn (primary) and Wind (src/wind_reconcile.py), used specifically
for verifying the most recent 1-2 months, since that's the window where
celma's own publishing lag (see HANDOFF.md 2026-08-11 entries) matters most
and the user doesn't have a fresh Wind export on hand every time.

## Why there's no listing/crawl automation here

Every provincial site's own *listing* page for this column was either
JS-rendered (empty on a plain fetch, same failure mode as celma's list-page
caching pitfall) or blocked by a WAF/cert mismatch when probed with
WebFetch/requests. There is no reliable "give me this province's bonds for
month X" entrypoint the way celma.org.cn's channelId-based pagination gives
us. So this module's unit of work is **one already-known announcement URL**
(found via WebSearch, same way PROVINCE_SOURCES below was populated) --
verify_announcement() parses it, diff_against_state() cross-checks it. This
mirrors wind_reconcile.py's own shape (a manually-supplied input, not a
live crawl).

## Why matching is (province, issue_date, term, amount), not bond_short_name

Confirmed empirically (2026-08-11) against real announcements from 江苏省
and 宁夏回族自治区: provincial announcements follow a nationally standardized
disclosure template (财库〔2020〕43号/36号) built around 债券名称/计划发行
规模/实际发行规模/发行期限/票面利率/付息.../到期日 -- this is a DIFFERENT
template from celma's own 表2-9/表2-10 aggregation table, and critically it
never carries the market 债券编码/债券简称 that wind_reconcile.py's matching
relies on. The only reliable join key across sources is therefore the
tuple (province, issue_date, term, amount) -- same fallback approach
wind_reconcile.py used internally to catch celma/Wind naming-convention
mismatches.

## Structure types (confirmed 2026-08-11, see HANDOFF.md for full findings)

- "pdf": cover HTML page links to a PDF attachment with the actual
  key-value bond blocks. Text-native PDFs work well; some provinces
  (confirmed: 江苏) serve *scanned* PDFs that only extract_pdf()'s OCR
  fallback can read, and OCR quality is materially worse -- expect missing
  fields (esp. 票面利率/发行期限), never fabricate them.
- "html": the bond blocks are inline in the cover page's own HTML, no
  attachment at all. Cleanest and most reliable source (confirmed: 宁夏,
  宁波) -- no OCR risk.
- "docx" (湖南): attachment is a Word doc containing an actual table (same
  field vocabulary as pdf/html, just already structured into cells) --
  parsed via parse_docx_table() using python-docx (now a project
  dependency, see requirements.txt). czt.hunan.gov.cn's TLS handshake also
  fails on this machine's OpenSSL (`BAD_ECPOINT`) via requests/urllib3;
  _fetch_html()/_fetch_bytes() fall back to a `curl` subprocess for it.
- "image" (天津): announcements are JPG scans, would need the same
  fitz+pytesseract OCR path pdf_extract.py uses for scanned PDFs, just
  applied to a standalone image instead of a PDF page. NOT IMPLEMENTED --
  calling verify_announcement() for it raises NotImplementedError rather
  than silently returning nothing.
- 重庆 (Chongqing): cover URL confirmed real but WebFetch couldn't extract
  a body during research -- structure genuinely unconfirmed, not just
  unimplemented. Don't assume "pdf" or "html" for it without checking.

Per explicit user instruction (2026-08-11): do NOT chase the remaining
~17 provinces where research only found partial leads or nothing at all
(甘肃/西藏/山西/陕西/山东/青岛/浙江/湖北/河南/江西/安徽/广东/吉林/黑龙江/
云南/四川/福建/内蒙古 -- 内蒙古 specifically confirmed to have NO
province-own results page at all, not just unfound). Only the provinces in
PROVINCE_SOURCES below are wired up; extend it when/if the user asks.
"""
import re
from dataclasses import dataclass
from pathlib import Path

from . import http_client, pdf_extract

AMOUNT_TOLERANCE_YI = 0.05
RATE_TOLERANCE_PCT = 0.011


@dataclass
class ProvinceSource:
    province: str
    domain: str
    structure: str  # "pdf" | "html" | "docx" | "image" | "unknown"
    example_url: str
    notes: str = ""


# One verified real 2026 announcement URL per province, from the 2026-08-11
# research pass (6 parallel agents, 27 provinces). example_url is not meant
# to be re-fetched forever -- it's the calibration sample structure was
# confirmed against; pass a fresh URL to verify_announcement() for new
# months.
PROVINCE_SOURCES: dict[str, ProvinceSource] = {
    "上海市": ProvinceSource(
        "上海市", "czj.sh.gov.cn", "pdf",
        "https://czj.sh.gov.cn/zys_8908/gsgg_8929/czywgg_8930/dfzwfxjg/20260713/xxfboswf0000003751.html",
    ),
    "新疆维吾尔自治区": ProvinceSource(
        "新疆维吾尔自治区", "czt.xinjiang.gov.cn", "pdf",
        "https://czt.xinjiang.gov.cn/xjczt/c115017/202606/dca280a6218e4b54be9e3c1a5c53c1b0.shtml",
    ),
    "河北省": ProvinceSource(
        "河北省", "czt.hebei.gov.cn", "pdf",
        "http://czt.hebei.gov.cn/root17/zfxx/202608/t20260804_2407370.html",
    ),
    "贵州省": ProvinceSource(
        "贵州省", "czt.guizhou.gov.cn", "pdf",
        "https://czt.guizhou.gov.cn/zwgk/zdlyxx/zfzw/202607/t20260709_90604174.html",
        notes="Publishes on a regular near-monthly cadence (批次二~五 found for 2026 alone) -- "
              "good candidate for a recurring monthly check.",
    ),
    "江苏省": ProvinceSource(
        "江苏省", "czt.jiangsu.gov.cn", "pdf",
        "https://czt.jiangsu.gov.cn/art/2026/7/24/art_77314_11808314.html",
        notes="Confirmed scanned/OCR PDF -- expect missing coupon_rate_pct/term on some rows. "
              "Sibling announcements (e.g. a same-day refinancing-bond batch) may exist at "
              "neighboring art_77314_<id>.html URLs not covered by a single fetch.",
    ),
    "宁夏回族自治区": ProvinceSource(
        "宁夏回族自治区", "czt.nx.gov.cn", "html",
        "https://czt.nx.gov.cn/xwzx/tzgg/202606/t20260610_5261427.html",
        notes="Cleanest source of the batch -- inline HTML, no PDF/OCR step, verified byte-exact "
              "against celma on all 6 bonds in the calibration sample.",
    ),
    "宁波市": ProvinceSource(
        "宁波市", "czj.ningbo.gov.cn", "html",
        "http://czj.ningbo.gov.cn/art/2025/5/26/art_1229029201_58890796.html",
        notes="2026-dated example not yet located (search-index lag) -- URL pattern "
              "art/YYYY/M/D/art_1229029201_<id>.html confirmed stable since 2023.",
    ),
    "天津市": ProvinceSource(
        "天津市", "cz.tj.gov.cn", "image",
        "https://cz.tj.gov.cn/zwgk_53713/zfzq/202608/t20260804_7347170.html",
        notes="Bond data is embedded as a sequence of JPG page-scan images (one per 'page' of the "
              "original document, filenames like W020260804628643438038_ORIGIN.jpg), not a single "
              "PDF/HTML doc -- parsed via parse_image_sequence() (downloads every content image "
              "filtered by the W<18+ digits>_ORIGIN.<ext> naming convention, OCRs each with "
              "pytesseract, concatenates page texts, reuses parse_announcement_text()). LOW YIELD "
              "in practice like 重庆 (Chongqing), but for a DIFFERENT reason: 天津's original table "
              "wraps the '计划发行规模'/'实际发行规模' cell labels onto two visual lines with other "
              "cells' text interposed between them, so pytesseract's line-by-line reading never sees "
              "'计划发行规模' as contiguous text -- amount/bond_name come back None on every row. "
              "Tested against the real 10-image, 2026-08-04 batch: only 2 of 10 rows resolved to an "
              "unambiguous single-candidate match (term+date alone, no amount to confirm), 7 correctly "
              "quarantined as low_confidence by diff_against_state's ambiguous-multi-candidate check "
              "(multiple 30Y bonds same day, nothing to disambiguate), 0 false matches or fabricated "
              "mismatches produced. Not worth chasing further -- the label-wrapping issue is a "
              "genuine OCR/layout limitation, not a parser bug.",
    ),
    "湖南省": ProvinceSource(
        "湖南省", "czt.hunan.gov.cn", "docx",
        "https://czt.hunan.gov.cn/czt/dzqzfzjxx/202606/t20260625_34011506.html",
        notes="Attachment is .docx, not .pdf -- but its content is an actual Word TABLE "
              "(债券名称/计划发行规模（亿元）/实际发行规模（亿元）/发行期限（年）/票面利率/...), "
              "cleaner than the regex-over-flattened-text path used for pdf/html since there's no "
              "block-splitting risk at all. czt.hunan.gov.cn's TLS handshake fails on this machine's "
              "OpenSSL (BAD_ECPOINT) via `requests`/urllib3 even though `curl` connects fine to the "
              "same URL -- _fetch_bytes() falls back to a `curl` subprocess on SSLError.",
    ),
    "重庆市": ProvinceSource(
        "重庆市", "czj.cq.gov.cn", "pdf",
        "https://czj.cq.gov.cn/zwgk_268/zfxxgkml/dfzfzw/202607/t20260729_15868828.html",
        notes="Structure confirmed 2026-08-11 via direct fetch (WebFetch had failed to surface the "
              "PDF link during research -- a direct requests.get()+regex found it fine, same "
              "P0<id>.pdf naming convention used by several other 政府信息公开 CMS instances). "
              "LOW YIELD in practice: this is a fully scanned PDF (no text layer, extract_pdf() "
              "always falls back to OCR) and its layout groups all field LABELS together before "
              "all their VALUES per bond (not the label-immediately-followed-by-value adjacency "
              "every other confirmed source uses), which the current regex parser doesn't handle -- "
              "tested against the real 6-bond 2026-07-29 batch and got 0 usable rows (most fields "
              "None, one row correctly caught by the low_confidence safety net). Safe (no false "
              "matches/mismatches were produced) but not currently useful for 重庆 specifically; "
              "would need a positional label/value pairing strategy instead of adjacency regex to "
              "improve, not attempted -- low ROI given how poor the source OCR itself is.",
    ),
}

_BOND_BLOCK_RE = re.compile(r"债券名称")
_NAME_RE = re.compile(r"^(.+?)计划发行规模")
_AMOUNT_RE = re.compile(r"实际发行规模([\d.]+)亿元")
_PLANNED_AMOUNT_RE = re.compile(r"计划发行规模([\d.]+)亿元")
# 宁波 confirmed (2026-08-11) to state amounts in 万元 (10k yuan), not 亿元
# like every other confirmed province -- e.g. "110000万元" == 11亿元. Only
# tried when the 亿元 patterns above miss, and converted by /10000.
_AMOUNT_WAN_RE = re.compile(r"实际发行规模([\d.]+)万元")
_PLANNED_AMOUNT_WAN_RE = re.compile(r"计划发行规模([\d.]+)万元")
_TERM_RE = re.compile(r"发行期限(\d+)年")
_RATE_RE = re.compile(r"票面利率([\d.]+)%")
_ISSUE_DATE_RE = re.compile(r"(\d{4})年(\d{1,2})月(\d{1,2})日已完成招标")
# 宁波 confirmed (2026-08-11) to never use the "已完成招标" phrasing at all
# -- the bid date only appears in the announcement's own title, e.g.
# "2025年5月23日宁波市政府债券发行结果公告". Tried as a fallback only.
_TITLE_DATE_RE = re.compile(r"(\d{4})年(\d{1,2})月(\d{1,2})日\S{0,4}(?:政府债券|地方政府债)发行结果公告")
# Some provinces' templates (confirmed: 新疆) also carry 债券简称 -- when
# present it's a far more reliable join key than (province,date,term), so
# capture it opportunistically. OCR is confirmed (新疆 sample) to misread
# "债" as "贵" here, so match either. Short names look like "26新疆债26":
# 2-digit year, then arbitrary CJK, then a 1-3 digit trailing number.
_SHORTNAME_RE = re.compile(r"[债贵]券简称(\d{2}[^\d]{1,10}\d{1,3})")


def _strip_html_to_text(html: str) -> str:
    text = re.sub(r"<script[^>]*>.*?</script>", " ", html, flags=re.S | re.I)
    text = re.sub(r"<style[^>]*>.*?</style>", " ", text, flags=re.S | re.I)
    text = re.sub(r"<[^>]+>", "|", text)
    import html as htmlmod
    return htmlmod.unescape(text)


def parse_announcement_text(raw_text: str, province: str) -> list[dict]:
    """Shared parser for the standardized 财库〔2020〕43号/36号 disclosure
    template -- confirmed identical field vocabulary whether the source is
    clean HTML (宁夏) or OCR'd PDF text (江苏/新疆), just varying in how much
    noise corrupts individual fields. Returns rows with `bond_name`,
    `bond_short_name` (when the template carries one -- not all do, see
    module docstring), `total_amount_yi`, `term`, `coupon_rate_pct`,
    `issue_date`, `province`.

    Splits on "债券名称" to find bond block boundaries. This is confirmed
    FRAGILE against heavy OCR noise -- two distinct failure modes seen in
    the 2026-08-11 calibration sample:
      1. 新疆 (2026-06-11, 3-bond batch): OCR drops/garbles the "债券名称"
         label on 2 of 3 bonds, split finds only 1 boundary, one bond's
         票面利率 ends up attached to a different bond's total_amount_yi.
         Caught by comparing 债券简称/债券代码 occurrences (present in
         richer templates like 新疆's) against the number of blocks found.
      2. 河北 (2026-08-03, 4-bond batch mixing 一般/专项/再融资): OCR is bad
         enough that even 债券简称 mentions are mostly dropped too (only 1
         of 4 survived), so check #1 alone doesn't fire. Caught instead by
         checking whether other field-label keywords (债券代码/存续期/发行
         期限/票面利率/付息/到期日/计划发行/实际发行) leaked into the
         `bond_name` capture itself -- a clean bond name (even a legitimate
         hyphenated dual-name like 宁夏's "A方案名-B方案名" for one bond)
         never contains those labels; their presence means _NAME_RE's
         non-greedy match ran past a garbled/missing block boundary and
         swallowed a second bond's title plus stray label text.
    Either signal marks every row from the document as low-confidence.
    Rows this uncertain must never be treated as a clean amount/rate
    mismatch by diff_against_state -- always inspect `warnings` first."""
    text = re.sub(r"\s+", "", raw_text).replace("|", "")

    date_m = _ISSUE_DATE_RE.search(text) or _TITLE_DATE_RE.search(text)
    issue_date = f"{date_m.group(1)}-{int(date_m.group(2)):02d}-{int(date_m.group(3)):02d}" if date_m else None

    blocks = _BOND_BLOCK_RE.split(text)[1:]
    expected_n = len(_SHORTNAME_RE.findall(text))
    shortname_count_mismatch = expected_n > 0 and expected_n != len(blocks)

    rows = []
    for b in blocks:
        name_m = _NAME_RE.match(b)
        amt_m = _AMOUNT_RE.search(b) or _PLANNED_AMOUNT_RE.search(b)
        amt_yi = float(amt_m.group(1)) if amt_m else None
        if amt_yi is None:
            wan_m = _AMOUNT_WAN_RE.search(b) or _PLANNED_AMOUNT_WAN_RE.search(b)
            if wan_m:
                amt_yi = round(float(wan_m.group(1)) / 10000, 4)
        term_m = _TERM_RE.search(b)
        rate_m = _RATE_RE.search(b)
        shortname_m = _SHORTNAME_RE.search(b)
        bond_name = name_m.group(1) if name_m else None
        leaked_labels = any(
            kw in (bond_name or "")
            for kw in ("债券代码", "存续期", "发行期限", "票面利率", "付息", "到期日", "计划发行", "实际发行")
        )
        row = {
            "province": province,
            "bond_name": bond_name,
            "bond_short_name": shortname_m.group(1) if shortname_m else None,
            "total_amount_yi": amt_yi,
            "term": f"{term_m.group(1)}Y" if term_m else None,
            "coupon_rate_pct": float(rate_m.group(1)) if rate_m else None,
            "issue_date": issue_date,
        }
        if shortname_count_mismatch:
            row["warnings"] = (
                f"文档内含{expected_n}处债券简称但只切分出{len(blocks)}个债券名称分块，"
                f"疑似OCR识别丢失部分'债券名称'标签导致相邻债券字段串块，"
                f"本行金额/期限/利率可能并非同一支债券，需人工核对原文"
            )
        elif leaked_labels:
            row["warnings"] = (
                f"债券名称字段内混入了'债券代码/存续期/票面利率'等字段标签文本，"
                f"疑似OCR丢失中间'债券名称'标签导致多支债券的标题和字段被合并进同一分块，"
                f"本行金额/期限/利率可能并非同一支债券，需人工核对原文"
            )
        rows.append(row)
    return rows


def _find_attachment_link(cover_html: str, base_url: str, ext: str) -> str | None:
    m = re.search(rf'href="([^"]+\.{ext})"', cover_html, re.I)
    if not m:
        return None
    href = m.group(1)
    if href.startswith("http"):
        return href
    from urllib.parse import urljoin
    return urljoin(base_url, href)


def _find_pdf_link(cover_html: str, base_url: str) -> str | None:
    return _find_attachment_link(cover_html, base_url, "pdf")


def _find_docx_link(cover_html: str, base_url: str) -> str | None:
    return _find_attachment_link(cover_html, base_url, "docx?")


def _fetch_bytes(url: str) -> bytes:
    """Like http_client.fetch_pdf() but for arbitrary binary attachments
    (currently just .docx), with a `curl` subprocess fallback. Confirmed
    (2026-08-11, czt.hunan.gov.cn) some .gov.cn TLS stacks trip this
    machine's OpenSSL 3.6.3 with a hard `[SSL: BAD_ECPOINT]` during the
    ECDHE handshake -- reproducible every time via `requests`/urllib3 and
    even bare `openssl s_client`, yet `curl` connects to the identical URL
    without issue (same class of incompatibility as the chinamoney.org.cn
    legacy-renegotiation problem found earlier this project -- different
    root cause, same "this machine's OpenSSL build vs. this server"
    shape). Not cached on disk like fetch_pdf()/extract_pdf() -- .docx
    announcements are infrequent enough that re-fetching each call is fine."""
    import requests
    headers = {"User-Agent": _config().USER_AGENT}
    try:
        resp = requests.get(url, headers=headers, timeout=20)
        resp.raise_for_status()
        return resp.content
    except requests.exceptions.SSLError:
        import subprocess
        result = subprocess.run(
            ["curl", "-sL", "-A", headers["User-Agent"], url],
            capture_output=True, timeout=45, check=True,
        )
        return result.stdout


def _config():
    from . import config
    return config


def _fetch_html(url: str, use_cache: bool = True) -> str:
    """Like http_client.fetch() but with the same curl fallback as
    _fetch_bytes() for the BAD_ECPOINT TLS failure -- needed because the
    cover-page HTML fetch hits the same affected hosts (confirmed:
    czt.hunan.gov.cn) as the .docx attachment fetch, not just the
    attachment step. Manually writes into http_client's own on-disk cache
    on the curl-fallback path too (http_client.fetch() itself never gets
    that far, since it raises before its own cache-write line runs) --
    without this, every call to an SSL-affected host reshells out to curl,
    which is slow enough (confirmed: occasionally exceeds a 30s timeout)
    that repeated calls in one session are worth avoiding."""
    try:
        return http_client.fetch(url, use_cache=use_cache)
    except RuntimeError as e:
        if "BAD_ECPOINT" not in str(e) and "SSLError" not in str(e):
            raise
        text = _fetch_bytes(url).decode("utf-8", errors="ignore")
        if use_cache:
            cache_file = http_client._cache_path(url, _config().RAW_HTML_DIR, ".html")
            try:
                cache_file.write_text(text, encoding="utf-8")
            except OSError:
                pass
        return text


def parse_docx_table(docx_bytes: bytes, province: str) -> list[dict]:
    """Parse a provincial 发行结果公告 .docx attachment. Confirmed (湖南,
    2026-06-25) these carry the same standardized field vocabulary as the
    pdf/html sources, but as an actual Word table -- read directly via
    column-header matching (same spirit as extract_result.py's
    _find_table_columns, simpler here since docx tables don't have the
    OCR/PDF layout-collapse failure modes that module was built around)."""
    import io
    import docx

    doc = docx.Document(io.BytesIO(docx_bytes))

    issue_date = None
    for p in doc.paragraphs:
        m = _ISSUE_DATE_RE.search(p.text.replace(" ", ""))
        if m:
            issue_date = f"{m.group(1)}-{int(m.group(2)):02d}-{int(m.group(3)):02d}"
            break

    field_map = {
        "债券名称": "bond_name", "实际发行规模": "total_amount_yi",
        "发行期限": "term", "票面利率": "coupon_rate_pct",
    }

    rows = []
    for table in doc.tables:
        if not table.rows:
            continue
        header = [c.text.strip() for c in table.rows[0].cells]
        col_idx = {}
        for i, h in enumerate(header):
            for cn_key, field_name in field_map.items():
                if cn_key in h and field_name not in col_idx.values():
                    col_idx[i] = field_name
        if "bond_name" not in col_idx.values():
            continue
        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if not any(cells):
                continue
            values = {}
            for i, field_name in col_idx.items():
                if i >= len(cells):
                    continue
                raw = cells[i]
                if field_name == "total_amount_yi":
                    m = re.search(r"[\d.]+", raw)
                    values[field_name] = float(m.group(0)) if m else None
                elif field_name == "term":
                    m = re.search(r"\d+", raw)
                    values[field_name] = f"{m.group(0)}Y" if m else None
                elif field_name == "coupon_rate_pct":
                    m = re.search(r"[\d.]+", raw)
                    values[field_name] = float(m.group(0)) if m else None
                else:
                    values[field_name] = raw or None
            rows.append({
                "province": province,
                "bond_name": values.get("bond_name"),
                "bond_short_name": None,
                "total_amount_yi": values.get("total_amount_yi"),
                "term": values.get("term"),
                "coupon_rate_pct": values.get("coupon_rate_pct"),
                "issue_date": issue_date,
            })
    return rows


# Content images on these sites (confirmed: 天津) follow a "W<18+ digits>_ORIGIN.<ext>"
# naming convention distinct from site-chrome assets (logos/icons/decorative
# images use short, human-readable filenames like "big.jpg"/"dy.jpg").
_CONTENT_IMAGE_RE = re.compile(r'src="([^"]*W\d{15,}_ORIGIN\.(?:jpg|jpeg|png))"', re.I)


def parse_image_sequence(cover_html: str, base_url: str, province: str, use_cache: bool = True) -> list[dict]:
    """Parse a provincial 发行结果公告 whose content is a sequence of scanned
    page images rather than a single PDF/HTML/docx document (confirmed: 天津).
    Downloads every content image referenced on the cover page, OCRs each
    with pytesseract (same chi_sim model pdf_extract.py uses for scanned
    PDFs), concatenates the page texts in DOM order -- mirroring how
    pdf_extract.py concatenates a multi-page PDF's text -- then reuses the
    same shared parse_announcement_text() every other source goes through,
    since the underlying field vocabulary is identical regardless of
    delivery mechanism."""
    if not pdf_extract.OCR_AVAILABLE:
        raise RuntimeError(
            "pytesseract/PyMuPDF/Pillow not available -- image OCR requires the same "
            "optional OCR dependencies pdf_extract.py uses for scanned PDFs"
        )
    import pytesseract
    from PIL import Image
    from urllib.parse import urljoin
    import io

    # Confirmed (天津, 2026-08-11): each content image's filename appears
    # TWICE in the page markup (a thumbnail <img> plus a full-size <a href>
    # pointing at the same file, or similar) -- de-dup by filename (not by
    # full URL, in case of a "./" vs bare-path prefix difference) or every
    # page gets OCR'd twice and every bond row comes back doubled.
    seen_filenames = set()
    image_urls = []
    for m in _CONTENT_IMAGE_RE.finditer(cover_html):
        filename = m.group(1).rsplit("/", 1)[-1]
        if filename in seen_filenames:
            continue
        seen_filenames.add(filename)
        image_urls.append(urljoin(base_url, m.group(1)))
    if not image_urls:
        raise RuntimeError(f"no content images found on cover page {base_url}")

    page_texts = []
    for img_url in image_urls:
        img_bytes = _fetch_bytes(img_url) if not use_cache else _fetch_bytes_cached(img_url)
        img = Image.open(io.BytesIO(img_bytes))
        page_texts.append(pytesseract.image_to_string(img, lang="chi_sim"))

    rows = parse_announcement_text("\n".join(page_texts), province)
    for r in rows:
        r["extraction_method"] = "ocr"
    return rows


def _fetch_bytes_cached(url: str) -> bytes:
    """Disk-cached wrapper around _fetch_bytes() for content images -- OCR is
    slow (confirmed: several seconds per page image), so repeated calls
    within/across sessions should hit the cache like pdf_extract.py's own
    .extract.json sidecar does, not re-download+re-OCR every time."""
    cache_file = http_client._cache_path(url, _config().RAW_PDF_DIR, Path(url).suffix or ".img")
    if cache_file.exists() and cache_file.stat().st_size > 0:
        return cache_file.read_bytes()
    data = _fetch_bytes(url)
    try:
        cache_file.write_bytes(data)
    except OSError:
        pass
    return data


def verify_announcement(province: str, url: str, use_cache: bool = True) -> list[dict]:
    """Fetch and parse one provincial 发行结果公告 URL. Returns a list of
    bond-row dicts (see parse_announcement_text). Raises NotImplementedError
    for structure types that aren't wired up yet (docx/image/unknown) --
    callers should catch this per-province rather than assuming every
    registered source is actually fetchable today."""
    src = PROVINCE_SOURCES.get(province)
    structure = src.structure if src else None

    if structure == "html":
        html = _fetch_html(url, use_cache=use_cache)
        text = _strip_html_to_text(html)
        return parse_announcement_text(text, province)

    if structure == "pdf":
        cover_html = _fetch_html(url, use_cache=use_cache)
        pdf_url = _find_pdf_link(cover_html, url)
        if not pdf_url:
            raise RuntimeError(f"no PDF link found on cover page {url}")
        pdf_path = http_client.fetch_pdf(pdf_url, use_cache=use_cache)
        result = pdf_extract.extract_pdf(pdf_path, use_cache=use_cache)
        rows = parse_announcement_text(result["text"], province)
        if result["method"] == "ocr":
            for r in rows:
                r["extraction_method"] = "ocr"
        return rows

    if structure == "docx":
        cover_html = _fetch_html(url, use_cache=use_cache)
        docx_url = _find_docx_link(cover_html, url)
        if not docx_url:
            raise RuntimeError(f"no .docx link found on cover page {url}")
        docx_bytes = _fetch_bytes(docx_url)
        return parse_docx_table(docx_bytes, province)

    if structure == "image":
        cover_html = _fetch_html(url, use_cache=use_cache)
        return parse_image_sequence(cover_html, url, province, use_cache=use_cache)

    raise NotImplementedError(
        f"structure '{structure}' for {province} is not implemented "
        f"(see PROVINCE_SOURCES notes) -- unknown needs a real structure check first"
    )


def diff_against_state(rows: list[dict], state_df) -> dict:
    """Cross-check parsed provincial rows against data/state_results.csv.
    Prefers matching on `bond_short_name` when the row has one (exact,
    reliable -- only some provincial templates carry it, see module
    docstring); falls back to (province, issue_date, term) otherwise, using
    amount to disambiguate same-day same-term multi-tranche bonds.

    Rows flagged `warnings` by parse_announcement_text (suspected
    OCR block-splitting cross-contamination) are routed to
    `low_confidence` instead of amount_mismatch/rate_mismatch -- a mismatch
    on a row whose fields might belong to two different bonds is not
    evidence of anything and must not be reported as a real discrepancy.

    Returns {"matched": [...], "amount_mismatch": [...], "rate_mismatch": [...],
    "not_found_in_state": [...], "low_confidence": [...]}, each a list of
    dicts with `row` (and `state_row` where matched), so a caller can print
    a clear report."""
    matched, amount_mismatch, rate_mismatch, not_found, low_confidence = [], [], [], [], []
    for row in rows:
        if row.get("warnings"):
            low_confidence.append({"row": row})
            continue

        candidates = None
        if row.get("bond_short_name"):
            hit = state_df[
                (state_df["province"] == row["province"])
                & (state_df["bond_short_name"] == row["bond_short_name"])
            ]
            if not hit.empty:
                candidates = hit

        if candidates is None:
            if row["issue_date"] is None or row["term"] is None:
                not_found.append({"row": row, "reason": "row itself missing issue_date/term, cannot match"})
                continue
            candidates = state_df[
                (state_df["province"] == row["province"])
                & (state_df["issue_date"] == row["issue_date"])
                & (state_df["term"] == row["term"])
            ]
            if candidates.empty:
                not_found.append({"row": row, "reason": "no state_results.csv row for this province/date/term"})
                continue
            if len(candidates) > 1:
                if row["total_amount_yi"] is not None:
                    close = candidates[(candidates["total_amount_yi"] - row["total_amount_yi"]).abs() <= AMOUNT_TOLERANCE_YI]
                    if len(close) == 1:
                        candidates = close
                if len(candidates) > 1:
                    # Multiple same-day same-term bonds and nothing (amount)
                    # to tell them apart -- picking candidates[0] here would
                    # be a coin flip dressed up as a "match". A row with no
                    # amount/rate at all (confirmed: 天津's OCR, which loses
                    # 计划/实际发行规模 to a table-layout quirk pytesseract
                    # can't recover) must never silently claim to match a
                    # SPECIFIC bond it can't actually distinguish.
                    low_confidence.append({
                        "row": row,
                        "reason": f"{len(candidates)}支{row['province']}债券在{row['issue_date']}同为{row['term']}期，"
                                  f"且本行缺少可用于区分的发行规模，无法确认具体对应哪一支，需人工核对原文",
                    })
                    continue

        srow = candidates.iloc[0]
        entry = {"row": row, "state_row": srow.to_dict()}
        if row["total_amount_yi"] is not None and abs(row["total_amount_yi"] - srow["total_amount_yi"]) > AMOUNT_TOLERANCE_YI:
            amount_mismatch.append(entry)
        elif row["coupon_rate_pct"] is not None and abs(row["coupon_rate_pct"] - srow["coupon_rate_pct"]) > RATE_TOLERANCE_PCT:
            rate_mismatch.append(entry)
        else:
            matched.append(entry)
    return {"matched": matched, "amount_mismatch": amount_mismatch, "rate_mismatch": rate_mismatch,
            "not_found_in_state": not_found, "low_confidence": low_confidence}
